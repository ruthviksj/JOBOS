"""Resume Agent (spec 25-26, 53).

Generates a profile-tailored resume (Markdown + DOCX) from the evidence library
and the selected profile. Anti-hallucination: every claim must be an evidence
record's allowed claim; factual fields (dates, companies, metrics) come from the
evidence library and are NOT invented or altered (spec 25).

DOCX generation via python-docx.
"""

from __future__ import annotations

from pathlib import Path

from sqlalchemy.orm import Session

from app import models

RESUME_DIR = Path(__file__).resolve().parent.parent.parent.parent / "RESUMES"


def _allowed_claims_for(db: Session, evidence_ids: list[str]) -> list[dict]:
    if not evidence_ids:
        return []
    ev = (
        db.query(models.Evidence)
        .filter(models.Evidence.evidence_id.in_(evidence_ids))
        .all()
    )
    out = []
    for e in ev:
        for claim in (e.allowed_claims or []):
            out.append({"claim": claim, "company": e.company, "strength": e.strength})
    return out


def build_resume_markdown(db: Session, profile: models.Profile,
                          evidence_ids: list[str]) -> str:
    claims = _allowed_claims_for(db, evidence_ids)
    lines = [
        f"# {profile.headline or profile.name}",
        "",
        profile.summary or profile.positioning_statement or "",
        "",
        "## Skills",
    ]
    for s in profile.skills or []:
        lines.append(f"- {s}")
    lines.append("")
    lines.append("## Experience (verified claims)")
    for c in claims:
        lines.append(f"- {c['claim']} — {c['company']} ({c['strength']})")
    if not claims:
        lines.append("- (no verified claims selected)")
    lines.append("")
    lines.append("## Domains")
    for d in profile.domains or []:
        lines.append(f"- {d}")
    return "\n".join(lines) + "\n"


def _render_docx(md: str, path: Path) -> None:
    from docx import Document

    doc = Document()
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(str(path))


def generate_resume(db: Session, profile_id: str,
                    evidence_ids: list[str] | None = None) -> dict:
    """Generate Markdown + DOCX resume for a profile. Returns saved paths."""
    profile = db.query(models.Profile).filter_by(profile_id=profile_id).first()
    if profile is None:
        raise ValueError(f"unknown profile {profile_id}")

    # Default evidence ids: the profile's recommended evidence set (spec 26).
    if not evidence_ids:
        evidence_ids = []
        for rule in profile.resume_rules or []:
            for tok in rule.split():
                if tok.startswith("EVID-"):
                    evidence_ids.append(tok)

    md = build_resume_markdown(db, profile, evidence_ids)
    RESUME_DIR.mkdir(parents=True, exist_ok=True)
    md_path = RESUME_DIR / f"{profile_id}-resume.md"
    docx_path = RESUME_DIR / f"{profile_id}-resume.docx"
    md_path.write_text(md, encoding="utf-8")
    _render_docx(md, docx_path)

    return {
        "profile_id": profile_id,
        "markdown": str(md_path),
        "docx": str(docx_path),
        "claims": len([c for c in _allowed_claims_for(db, evidence_ids)]),
    }
